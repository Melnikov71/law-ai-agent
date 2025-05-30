from dataclasses import dataclass
from datetime import date
import os
from pathlib import Path
import re
from typing import Sequence
from uuid import uuid4

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from dotenv import find_dotenv, load_dotenv
import img2pdf
from langchain_core.language_models import LanguageModelLike
from langchain_core.runnables import RunnableConfig
from langchain_core.tools import BaseTool, tool
from langchain_gigachat.chat_models import GigaChat
from langgraph.prebuilt import create_react_agent
from langgraph.checkpoint.memory import InMemorySaver

load_dotenv(find_dotenv())

INCOMING_FILE = "Приглашение.jpg"
AGREEMENT_TEMPLATE = "agreement_template.docx"


@dataclass
class IncomingDocument:
    doc_number: str
    doc_date: date


@dataclass
class Products:
    okpd: str
    placement: str


class AgreementGenerator:
    """Класс генерации документа-согласия на участие в аукционе"""

    def __init__(self, incoming_document_data: IncomingDocument, products: list[Products]):
        self.template = AGREEMENT_TEMPLATE
        if not os.path.exists(self.template):
            raise FileNotFoundError(f"Шаблон '{self.template}' не найден в текущей директории.")

        self.incoming_document_data = incoming_document_data
        self.products = products
        self.output_document = f'Согласие на участие в аукционе по приглашению № {incoming_document_data.doc_number} от {incoming_document_data.doc_date.strftime("%d.%m.%Y")}.docx'

    @staticmethod
    def _replace_placeholder(paragraph, placeholder, replacement):
        # Заменяем плейсхолдеры в параграфах
        for run in paragraph.runs:
            if placeholder in run.text:
                run.text = run.text.replace(placeholder, replacement)

    @staticmethod
    def _sanitize_filename(filename):
        """
        Удаляет недопустимые символы из имени файла.
        """
        # Заменяем недопустимые символы на "_"
        return re.sub(r'[<>:"/\\|?*]', '_', filename)

    def generate(self):
        # Открываем шаблон
        doc = Document(self.template)

        # Проставляем исходящий номер и дату приглашения
        for paragraph in doc.paragraphs:
            self._replace_placeholder(paragraph, 'incomingnumber', self.incoming_document_data.doc_number)
            self._replace_placeholder(paragraph, 'incomingdate',
                                      self.incoming_document_data.doc_date.strftime('%d.%m.%Y'))

        # Заполняем таблицу с продукцией
        table = None
        for t in doc.tables:
            if 'Класс ОКПД' in t.cell(0, 0).text and 'Территориальное размещение (федеральный округ)' in t.cell(0,
                                                                                                                1).text:
                table = t
                break

        if table:
            # Очищаем существующие строки, кроме заголовка
            while len(table.rows) > 1:
                table._tbl.remove(table.rows[-1]._tr)

            # Добавляем новые строки с данными
            for product in self.products:
                row_cells = table.add_row().cells
                row_cells[0].text = product.okpd
                row_cells[1].text = product.placement

                # Выравниваем текст по центру
                for cell in row_cells:
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Добавляем рамку к ячейкам
                for cell in row_cells:
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    tcBorders = tcPr.first_child_found_in("w:tcBorders")
                    if tcBorders is None:
                        tcBorders = OxmlElement('w:tcBorders')
                        tcPr.append(tcBorders)

                    for border_name in ('top', 'left', 'bottom', 'right'):
                        border = OxmlElement(f'w:{border_name}')
                        border.set(qn('w:val'), 'single')
                        border.set(qn('w:sz'), '4')  # Толщина линии
                        border.set(qn('w:space'), '0')
                        border.set(qn('w:color'), '000000')  # Цвет линии
                        tcBorders.append(border)

        # Сохраняем заполненный документ
        doc.save(self._sanitize_filename(self.output_document))


class LLMAgent:
    def __init__(self, model: LanguageModelLike, tools: Sequence[BaseTool]):
        self._model = model
        self._agent = create_react_agent(
            model,
            tools=tools,
            checkpointer=InMemorySaver())
        self._config: RunnableConfig = {
            "configurable": {"thread_id": uuid4().hex}}

    def upload_file(self, file):
        file_uploaded_id = self._model.upload_file(file).id_  # type: ignore
        return file_uploaded_id

    def invoke(
            self,
            content: str,
            attachments: list[str] | None = None,
            temperature: float = 1.0
    ) -> str:
        """Отправляет сообщение в чат"""
        message: dict = {
            "role": "user",
            "content": content,
            **({"attachments": attachments} if attachments else {})
        }
        return self._agent.invoke(
            {
                "messages": [message],
                "temperature": temperature
            },
            config=self._config)["messages"][-1].content


@tool
def generate_request_document(incoming_document: IncomingDocument, products: list[Products]) -> None:
    """Генерирует документ-заявку на участие в аукционе
    Arg:
        incoming_document: Номер и дата документа, на основании которого создаётся заявка
        products: Список продукции
    """
    AgreementGenerator(incoming_document, products).generate()


def print_agent_response(llm_response: str) -> None:
    print(f"\033[35m{llm_response}\033[0m")


def get_user_prompt() -> str:
    return input("\nТы: ")


def main():
    model = GigaChat(
        # model="GigaChat-2-Max",
        model="GigaChat-Pro",
        verify_ssl_certs=False,
    )

    agent = LLMAgent(model, tools=[generate_request_document])
    # system_prompt = (
    #     "Твоя задача взять из исходного файла исходящий номер, дату документа в формате день.месяц.год и"
    #     "список продукции из таблицы в исходном документе. Посоле этого на основании загруженного файла шаблона "
    #     "сгенерируй документ-согласие на участие в аукционе. Никакие данные не придумывай, всё необходимое строго "
    #     "возьми из загруженного файла."
    # )
    system_prompt = ("Твоя задача извлечь исходящий номер, дату документа и список видов продукции с полями 'Класс "
                     "ОКПД' и 'Территориальное размещение (федеральный округ)' из предоставленного документа. "
                     "Не придумывай никаких данных, всё необходимое возьми из предоставленного документа.")

    output_file = INCOMING_FILE

    if Path(INCOMING_FILE).suffix.lower() in [".jpg", ".jpeg"]:
        # Конвертируем изображение в PDF
        output_file = Path(INCOMING_FILE).with_suffix(".pdf")
        with open(output_file, "wb") as f:
            f.write(img2pdf.convert(INCOMING_FILE))

    file_uploaded_id = agent.upload_file(open(output_file, "rb"))

    agent_response = agent.invoke(content=system_prompt, attachments=[file_uploaded_id])

    if output_file != INCOMING_FILE:
        os.remove(output_file)

    while (True):
        print_agent_response(agent_response)
        agent_response = agent.invoke(get_user_prompt())


if __name__ == "__main__":
    try:
        main()
    except KeyboardInterrupt:
        print("\nдI`ll be back!")
